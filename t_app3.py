# ==============================================================================
# === Whisper & Gemini Transcriber App (v4.4 - No Console Windows) ===
# ==============================================================================
#
# Финальная, полностью рабочая версия.
# - Полностью убраны ВСЕ всплывающие консольные окна при вызове FFmpeg/FFprobe.
# - Восстановлен полный промт для Gemini.
# - Сохранение и загрузка API ключа Gemini из файла.
# - Улучшенный UX/UI.
#
# Зависимости:
# pip install torch transformers datasets pydub google-generativeai python-docx
# ==============================================================================

# --- Блок импорта библиотек ---
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox, Toplevel
import threading
import os
import sys
import math
import time
import tempfile
import subprocess
import json # Для парсинга вывода ffprobe

# Библиотеки для работы с аудио и ИИ
from transformers import pipeline
from pydub import AudioSegment
from pydub.utils import mediainfo # Импортируем для патчинга
from pydub.exceptions import CouldntDecodeError

# Библиотеки для анализа и экспорта в Word
import google.generativeai as genai
from docx import Document
from docx.shared import Inches

# --- СКРЫТИЕ КОНСОЛИ ДЛЯ SUBPROCESS В WINDOWS ---
if sys.platform == "win32":
    # Создаем глобальный объект startupinfo, который будет скрывать окно
    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    # Этого флага нет в стандартном модуле, но он нужен для Popen
    subprocess.CREATE_NO_WINDOW = 0x08000000

    # Патчим Popen, чтобы все вызовы использовали startupinfo
    _old_popen = subprocess.Popen
    def _new_popen(*args, **kwargs):
        kwargs['startupinfo'] = startupinfo
        kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
        return _old_popen(*args, **kwargs)
    subprocess.Popen = _new_popen
    
    ### НОВЫЙ КОД: Патчим внутренний метод pydub, который вызывает ffprobe ###
    def _new_get_media_info(filepath, timeout=None):
        """
        Замена для pydub.utils.mediainfo, которая вызывает ffprobe без окна.
        """
        cmd = [
            AudioSegment.ffprobe, 
            "-v", "quiet", 
            "-print_format", "json", 
            "-show_format", "-show_streams", 
            filepath
        ]
        
        # Используем Popen напрямую, чтобы наш патч сработал
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        
        # Ждем завершения и получаем вывод
        stdout, stderr = proc.communicate(timeout=timeout)
        
        # Проверяем, была ли ошибка
        if proc.returncode != 0:
            raise OSError(f"ffprobe error: {stderr.decode('utf-8')}")
            
        return json.loads(stdout)

    # Заменяем старую функцию на нашу новую
    mediainfo.get_media_info = _new_get_media_info


# --- ВСПОМОГАТЕЛЬНАЯ ФУНКЦИЯ ДЛЯ ОПРЕДЕЛЕНИЯ ПУТИ ---
# ... (остается без изменений) ...
def get_resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        base_path = os.path.dirname(sys.executable)
    else:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)

# --- КОНСТАНТЫ И ПРОМПТЫ (без изменений) ---
# ... (весь блок без изменений) ...
AVAILABLE_MODELS = { "Крошечная (tiny)": "openai/whisper-tiny", "Базовая (base)": "openai/whisper-base", "Малая (small)": "openai/whisper-small", "Средняя (medium)": "openai/whisper-medium", "Большая (large-v3)": "openai/whisper-large-v3", }
CHUNK_THRESHOLD_MS = 2 * 60 * 1000
CHUNK_LENGTH_MS = 1 * 60 * 1000
DEFAULT_GEMINI_PROMPT = """
# РОЛЬ
Ты — экспертный редактор и старший аналитик. Твоя задача — преобразовать сырую, неструктурированную и полную ошибок автоматическую расшифровку аудиозаписи в профессиональный, многоуровневый документ, готовый для вставки в Microsoft Word.

# КОНТЕКСТ
Ниже приведена сырая текстовая расшифровка аудиофайла. Текст содержит множество ошибок распознавания, особенно в технических терминах, названиях и именах собственных. В нем отсутствуют знаки препинания, абзацы и какая-либо структура. Это может быть монолог, диалог или групповое обсуждение.

# ОСНОВНАЯ ЗАДАЧА: Редактирование и коррекция (Два уровня)

**Уровень 1: Отредактированная версия (для чтения)**
Создай чистую и легко читаемую версию текста. Для этого:
1.  **Исправь все ошибки:** Удели особое внимание техническим терминам, именам собственным (компании, люди, технологии), омофонам и похожим по звучанию словам.
2.  **Приведи в порядок грамматику:** Расставь все знаки препинания, исправь орфографические и грамматические ошибки.
3.  **Улучши читаемость:** Разбей текст на логические абзацы. Удали мешающие слова-паразиты и заполнители пауз (эээ, ну, как бы, это самое), если они не несут важной смысловой нагрузки.
4.  **Оформи диалог:** Если из контекста очевиден диалог, оформи его репликами, используя тире (—).
5.  **Текст должен быть структирован и разбит на абзацы для удобства**

**Уровень 2: Дословная версия (для архива)**
Создай вторую, дословную версию текста. Для этого:
1.  **Исправь только ошибки распознавания:** Исправь неправильно распознанные слова на верные (например, "нераморфный" -> "нейроморфный").
2.  **Сохрани оригинальную речь:** **Не удаляй** слова-паразиты, повторы, незаконченные фразы и разговорные конструкции. Сохрани их, чтобы передать живую речь максимально точно.
3.  **Добавь базовую пунктуацию:** Расставь только точки в конце предложений и заглавные буквы, чтобы текст был минимально читаемым, но оставался "сырым" по стилю. Разбей на абзацы по смене говорящего или темы.
4.  Предоставь результат в виде списка сегментов, где каждый сегмент содержит время начала (start), время конца (end) и текст (text). Выведи результат в формате, где каждая строка представляет один сегмент:
[<start>s - <end>s] <text>


# ВТОРИЧНАЯ ЗАДАЧА: Анализ и извлечение сути

После создания двух версий текста выполни его глубокий анализ:
1.  **Краткое содержание (Executive Summary):** В 3-5 абцах изложи основную суть всего разговора.
2.  **Ключевые тезисы:** Представь в виде маркированного списка (5-7 пунктов) главные мысли, выводы и решения.
3.  **Словарь терминов:** Составь таблицу с ключевыми техническими понятиями и их кратким объяснением на основе контекста.
4.  **Вопросы и задачи:** Если в тексте обсуждались конкретные задачи, нерешенные вопросы или планы на будущее, составь их отдельным списком.

# ФОРМАТИРОВАНИЕ ДЛЯ WORD
Для обеспечения идеальной вставки в Microsoft Word, используй следующее Markdown-форматирование. Оно будет корректно преобразовано в стили Word:
*   `# Заголовок первого уровня` (станет "Заголовок 1" в Word).
*   `## Заголовок второго уровня` (станет "Заголовок 2" в Word).
*   `* Элемент списка` (станет маркированным списком).
*   `**Жирный текст**` (станет полужирным).
*   `— Реплика диалога` (станет текстом с тире).
*   Таблицы создавай с помощью `| Столбец 1 | Столбец 2 |`.

# СТРУКТУРА ВЫВОДА
Предоставь финальный результат в строгом соответствии со следующей структурой:

# Аналитический отчет по расшифровке аудио

## 1. Краткое содержание (Executive Summary)
[Здесь должно быть саммари]

## 2. Ключевые тезисы
[Здесь должен быть список тезисов]

## 3. Словарь ключевых терминов
[Здесь должна быть таблица с терминами]

## 4. Обсуждавшиеся вопросы и задачи
[Здесь должен быть список вопросов и задач]

## 5. Отредактированная версия текста (для публикации и чтения)
[Здесь должен быть полностью исправленный, "причесанный" и легко читаемый текст]

## 6. Дословная расшифровка (для архива и детального анализа)
[Здесь должен быть дословный, но исправленный по орфографии текст с сохранением всех особенностей живой речи]


---
Приступай к работе. Вот текст для обработки:
"""

# ==============================================================================
# === КЛАСС ПРИЛОЖЕНИЯ (GUI на Tkinter и TTK) ===
# ==============================================================================

class WhisperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Whisper & Gemini Transcriber v4.4")
        self.root.geometry("900x750")
        self.root.minsize(700, 500)

        # Блок инициализации FFmpeg
        try:
            ffmpeg_path = get_resource_path("ffmpeg.exe")
            ffprobe_path = get_resource_path("ffprobe.exe")
            if not os.path.exists(ffmpeg_path) or not os.path.exists(ffprobe_path):
                 raise FileNotFoundError("Не найден ffmpeg.exe или ffprobe.exe")
            AudioSegment.converter = ffmpeg_path
            AudioSegment.ffprobe = ffprobe_path
        except Exception as e:
            messagebox.showerror(
                "Критическая ошибка: FFmpeg не найден",
                f"Не удалось найти FFmpeg. Убедитесь, что 'ffmpeg.exe' и 'ffprobe.exe' находятся в той же папке, что и приложение.\n\nПуть поиска: {get_resource_path('')}\nОшибка: {e}"
            )
            self.root.after(100, self.root.destroy)
            return

        # ... (остальной код __init__ и весь класс без изменений) ...
        self.models_cache = {}
        self.current_model_name = tk.StringVar(value=list(AVAILABLE_MODELS.keys())[1])
        self.output_folder = tk.StringVar(value=os.getcwd())
        self.show_timestamps = tk.BooleanVar(value=True)
        self.stop_processing = threading.Event()
        self.gemini_api_key = None
        self.api_key_path = get_resource_path("gemini_api_key.txt")
        self.last_clean_transcription = ""
        style = ttk.Style(self.root)
        style.theme_use('clam')
        self._create_context_menu()
        self.create_widgets()
    def _create_context_menu(self):
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Копировать", command=self._copy_selection)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Выделить всё", command=self._select_all)
    def _show_context_menu(self, event):
        self.context_menu.tk_popup(event.x_root, event.y_root)
    def _copy_selection(self):
        try:
            widget = self.root.focus_get()
            selected_text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)
        except tk.TclError: pass
    def _select_all(self):
        try:
            widget = self.root.focus_get()
            widget.tag_add(tk.SEL, "1.0", tk.END)
            return 'break'
        except tk.TclError: pass
    def _custom_ask_string(self, title, prompt, parent):
        dialog = Toplevel(parent)
        dialog.title(title)
        dialog.transient(parent)
        dialog.grab_set()
        dialog.resizable(False, False)
        result = None
        entry_var = tk.StringVar()
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill="both", expand=True)
        ttk.Label(main_frame, text=prompt, wraplength=300).pack(pady=(0, 10))
        entry = ttk.Entry(main_frame, textvariable=entry_var, width=50)
        entry.pack(pady=5)
        entry_menu = tk.Menu(dialog, tearoff=0)
        def paste_from_clipboard():
            try: entry.insert(tk.INSERT, dialog.clipboard_get())
            except tk.TclError: pass
        entry_menu.add_command(label="Вставить", command=paste_from_clipboard)
        entry.bind("<Button-3>", lambda e: entry_menu.tk_popup(e.x_root, e.y_root))
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=(10, 0))
        def on_ok():
            nonlocal result
            result = entry_var.get()
            dialog.destroy()
        def on_cancel():
            nonlocal result
            result = None
            dialog.destroy()
        ttk.Button(button_frame, text="OK", command=on_ok, width=10).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Cancel", command=on_cancel, width=10).pack(side="left", padx=5)
        dialog.protocol("WM_DELETE_WINDOW", on_cancel)
        entry.focus_set()
        dialog.wait_window()
        return result
    def _load_api_key_from_file(self):
        try:
            if os.path.exists(self.api_key_path):
                with open(self.api_key_path, 'r', encoding='utf-8') as f:
                    key = f.read().strip()
                    if key:
                        self.gemini_api_key = key
                        self.log(f"API ключ Gemini загружен из файла.", tag='SUCCESS')
                        return True
        except Exception as e:
            self.log(f"Ошибка чтения файла с API ключом: {e}", tag='ERROR')
        return False
    def _save_api_key_to_file(self, key):
        try:
            with open(self.api_key_path, 'w', encoding='utf-8') as f:
                f.write(key)
            self.log(f"API ключ сохранен в {self.api_key_path}", tag='SUCCESS')
        except Exception as e:
            self.log(f"Не удалось сохранить API ключ в файл: {e}", tag='ERROR')
    def create_widgets(self):
        controls_frame = ttk.Frame(self.root, padding="10")
        controls_frame.pack(fill=tk.X, side=tk.TOP)
        ttk.Label(controls_frame, text="Модель Whisper:").pack(side=tk.LEFT, padx=(0, 5))
        self.model_combobox = ttk.Combobox(controls_frame, textvariable=self.current_model_name, values=list(AVAILABLE_MODELS.keys()), state="readonly", width=18)
        self.model_combobox.pack(side=tk.LEFT, padx=5)
        self.select_files_button = ttk.Button(controls_frame, text="Выбрать аудиофайлы...", command=self.select_audio_files)
        self.select_files_button.pack(side=tk.LEFT, padx=5)
        self.select_folder_button = ttk.Button(controls_frame, text="Папка для результатов...", command=self.select_output_folder)
        self.select_folder_button.pack(side=tk.LEFT, padx=5)
        self.timestamps_check = ttk.Checkbutton(controls_frame, text="Показывать таймстампы", variable=self.show_timestamps)
        self.timestamps_check.pack(side=tk.LEFT, padx=10)
        files_frame = ttk.Labelframe(self.root, text="Файлы для обработки", padding="10")
        files_frame.pack(fill=tk.BOTH, expand=False, padx=10, pady=5)
        scrollbar = ttk.Scrollbar(files_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.files_listbox = tk.Listbox(files_frame, selectmode=tk.EXTENDED, yscrollcommand=scrollbar.set, height=5)
        self.files_listbox.pack(fill=tk.X)
        scrollbar.config(command=self.files_listbox.yview)
        progress_frame = ttk.Frame(self.root, padding="10")
        progress_frame.pack(fill=tk.X, side=tk.TOP)
        progress_frame.columnconfigure(0, weight=1)
        self.progressbar = ttk.Progressbar(progress_frame, mode='determinate')
        self.progressbar.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.run_button = ttk.Button(progress_frame, text="Начать расшифровку", command=self.start_transcription_thread, state="disabled")
        self.run_button.grid(row=0, column=1, padx=5)
        self.cancel_button = ttk.Button(progress_frame, text="Отмена", command=self.stop_thread, state="disabled")
        self.cancel_button.grid(row=0, column=2)
        log_frame = ttk.Labelframe(self.root, text="Лог выполнения и результаты", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        self.log_area = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, font=("Segoe UI", 9))
        self.log_area.pack(fill=tk.BOTH, expand=True)
        self.log_area.bind("<Key>", lambda e: "break")
        self.log_area.bind("<Button-3>", self._show_context_menu)
        self.log_area.tag_config('INFO', foreground='blue')
        self.log_area.tag_config('SUCCESS', foreground='green')
        self.log_area.tag_config('ERROR', foreground='red')
        self.log_area.tag_config('HEADER', foreground='navy', font=("Segoe UI", 10, "bold"))
        self.log_area.tag_config('TIMESTAMP', foreground='gray')
        bottom_frame = ttk.Frame(self.root, padding="10")
        bottom_frame.pack(fill=tk.X, side=tk.BOTTOM)
        self.copy_button = ttk.Button(bottom_frame, text="Копировать весь лог", command=self.copy_log_to_clipboard)
        self.copy_button.pack(side=tk.LEFT)
        self.save_button = ttk.Button(bottom_frame, text="Сохранить лог в файл...", command=self.save_log_to_file)
        self.save_button.pack(side=tk.LEFT, padx=10)
        self.gemini_button = ttk.Button(bottom_frame, text="Анализировать с Gemini...", command=self.open_gemini_analyzer_window)
        self.gemini_button.pack(side=tk.RIGHT, padx=10)
    def log(self, message, tag=None):
        self.log_area.insert(tk.END, str(message) + "\n", tag)
        self.log_area.see(tk.END)
        self.root.update_idletasks()
    def select_audio_files(self):
        paths = filedialog.askopenfilenames(title="Выберите один или несколько аудиофайлов", filetypes=[("Аудиофайлы", "*.wav *.mp3 *.m4a *.ogg *.flac"), ("Все файлы", "*.*")])
        if paths:
            self.files_listbox.delete(0, tk.END)
            for path in paths:
                self.files_listbox.insert(tk.END, path)
            self.run_button.config(state="normal")
    def select_output_folder(self):
        path = filedialog.askdirectory(title="Выберите папку для сохранения результатов")
        if path:
            self.output_folder.set(path)
            self.log(f"Папка для сохранения результатов: {path}", tag='INFO')
    def toggle_ui_state(self, enabled=True):
        state = "normal" if enabled else "disabled"
        self.run_button.config(state=state)
        self.cancel_button.config(state="disabled" if enabled else "normal")
        self.select_files_button.config(state=state)
        self.select_folder_button.config(state=state)
        self.model_combobox.config(state="readonly" if enabled else "disabled")
        self.timestamps_check.config(state=state)
        self.copy_button.config(state=state)
        self.save_button.config(state=state)
        self.gemini_button.config(state=state)
    def stop_thread(self):
        self.log("Получен сигнал отмены. Завершение после текущей операции...", tag='ERROR')
        self.stop_processing.set()
    def copy_log_to_clipboard(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.log_area.get(1.0, tk.END))
        self.log("Лог скопирован в буфер обмена.", tag='SUCCESS')
    def save_log_to_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Текстовые файлы", "*.txt"), ("Все файлы", "*.*")], title="Сохранить лог как...")
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.log_area.get(1.0, tk.END))
                self.log(f"Лог успешно сохранен в: {file_path}", tag='SUCCESS')
            except Exception as e:
                self.log(f"Ошибка сохранения лога: {e}", tag='ERROR')
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")
    def start_transcription_thread(self):
        if self.files_listbox.size() == 0:
            messagebox.showwarning("Нет файлов", "Пожалуйста, выберите хотя бы один аудиофайл.")
            return
        self.toggle_ui_state(enabled=False)
        self.stop_processing.clear()
        self.log_area.delete(1.0, tk.END)
        self.last_clean_transcription = ""
        threading.Thread(target=self.process_files, daemon=True).start()
    def process_files(self):
        start_time = time.time()
        all_files_transcription = []
        try:
            selected_model_display_name = self.current_model_name.get()
            selected_model_id = AVAILABLE_MODELS[selected_model_display_name]
            if selected_model_id not in self.models_cache:
                self.log(f"Загрузка модели '{selected_model_display_name}'...", tag='INFO')
                self.models_cache[selected_model_id] = pipeline("automatic-speech-recognition", model=selected_model_id)
                self.log("Модель успешно загружена.", tag='SUCCESS')
            whisper_model = self.models_cache[selected_model_id]
            files_to_process = self.files_listbox.get(0, tk.END)
            self.progressbar['value'] = 0
            self.progressbar['maximum'] = len(files_to_process)
            for i, file_path in enumerate(files_to_process):
                if self.stop_processing.is_set(): break
                self.log(f"\n--- Обработка файла {i+1}/{len(files_to_process)}: {os.path.basename(file_path)} ---", tag='HEADER')
                try:
                    sound = AudioSegment.from_file(file_path)
                    sound = sound.set_channels(1).set_frame_rate(16000)
                    result = self.transcribe_segment(sound, whisper_model)
                    clean_text = result["text"].strip()
                    all_files_transcription.append(clean_text)
                    output_filename = os.path.splitext(os.path.basename(file_path))[0] + ".txt"
                    output_path = os.path.join(self.output_folder.get(), output_filename)
                    with open(output_path, 'w', encoding='utf-8') as f: f.write(clean_text)
                    self.log(f"Результат сохранен в: {output_path}", tag='SUCCESS')
                except CouldntDecodeError as e: self.log(f"ОШИБКА: Не удалось декодировать файл. Убедитесь, что FFmpeg найден и доступен. Ошибка: {e}", tag='ERROR')
                except Exception as e: self.log(f"ОШИБКА при обработке файла: {e}", tag='ERROR')
                finally: self.progressbar['value'] = i + 1
            self.last_clean_transcription = "\n\n---\n\n".join(all_files_transcription)
            total_elapsed_time = time.time() - start_time
            self.log(f"\n--- Пакетная обработка завершена за {total_elapsed_time:.2f} сек ---", tag='SUCCESS')
            if not self.stop_processing.is_set(): messagebox.showinfo("Готово", f"Обработка {len(files_to_process)} файлов завершена.")
        except Exception as e:
            self.log(f"\nКРИТИЧЕСКАЯ ОШИБКА: {e}", tag='ERROR')
            messagebox.showerror("Критическая ошибка", f"Произошла непредвиденная ошибка:\n{e}")
        finally: self.toggle_ui_state(enabled=True)
    def transcribe_segment(self, sound, model):
        start_time = time.time()
        full_transcription_result = {"text": "", "chunks": []}
        self.log(f"Аудио подготовлено. Длительность: {len(sound)/1000.0:.2f} сек.", tag='INFO')
        temp_dir = tempfile.gettempdir()
        if len(sound) > CHUNK_THRESHOLD_MS:
            self.log(f"Файл длиннее {CHUNK_THRESHOLD_MS // 1000 // 60} минут. Начинаю нарезку...", tag='INFO')
            num_chunks = math.ceil(len(sound) / CHUNK_LENGTH_MS)
            total_duration_offset = 0.0
            for i in range(num_chunks):
                if self.stop_processing.is_set(): return full_transcription_result
                start_ms = i * CHUNK_LENGTH_MS
                end_ms = (i + 1) * CHUNK_LENGTH_MS
                self.log(f"  Обработка куска {i+1}/{num_chunks}...", tag='INFO')
                chunk = sound[start_ms:end_ms]
                chunk_path = os.path.join(temp_dir, f"_temp_whisper_chunk_{i}.wav")
                chunk.export(chunk_path, format="wav")
                result = model(chunk_path, return_timestamps=True)
                os.remove(chunk_path)
                if "chunks" in result:
                    for segment in result["chunks"]:
                        if segment['timestamp'][0] is not None and segment['timestamp'][1] is not None:
                            segment['timestamp'] = (segment['timestamp'][0] + total_duration_offset, segment['timestamp'][1] + total_duration_offset)
                        full_transcription_result["chunks"].append(segment)
                total_duration_offset += len(chunk) / 1000.0
                elapsed_time = time.time() - start_time
                self.log(f"  Кусок {i+1} обработан. Прошло времени: {elapsed_time:.2f} сек.", tag='TIMESTAMP')
            full_transcription_result["text"] = " ".join([c['text'].strip() for c in full_transcription_result["chunks"]])
            self.log("Все куски обработаны и сшиты.", tag='SUCCESS')
        else:
            self.log(f"Файл короче {CHUNK_THRESHOLD_MS // 1000 // 60} минут. Обрабатываю целиком...", tag='INFO')
            temp_wav_path = os.path.join(temp_dir, "_temp_whisper_audio.wav")
            sound.export(temp_wav_path, format="wav")
            full_transcription_result = model(temp_wav_path, return_timestamps=True)
            os.remove(temp_wav_path)
        final_text = full_transcription_result["text"].strip()
        self.log("\n--- Полный текст ---", tag='HEADER')
        self.log(final_text if final_text else "[Тишина или нераспознанная речь]")
        if self.show_timestamps.get():
            self.log("\n--- Сегменты с таймстампами ---", tag='HEADER')
            if "chunks" in full_transcription_result and full_transcription_result["chunks"]:
                 for segment in full_transcription_result["chunks"]:
                    start, end = segment['timestamp']
                    if start is not None and end is not None:
                        self.log(f"[{start:.2f}s - {end:.2f}s]{segment['text']}", tag='TIMESTAMP')
            else:
                self.log("Сегменты не доступны.", tag='INFO')
        return full_transcription_result
    def open_gemini_analyzer_window(self):
        if not self.last_clean_transcription:
            messagebox.showwarning("Нет данных", "Сначала выполните расшифровку. Текст для анализа не найден.")
            return
        analyzer_window = Toplevel(self.root)
        analyzer_window.title("Анализ текста с помощью Gemini")
        analyzer_window.geometry("800x600")
        analyzer_window.transient(self.root)
        analyzer_window.grab_set()
        main_frame = ttk.Frame(analyzer_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        self._load_api_key_from_file()
        top_panel = ttk.Frame(main_frame)
        top_panel.pack(fill=tk.X, pady=(0, 10))
        def use_default(): prompt_text.delete(1.0, tk.END); prompt_text.insert(tk.END, DEFAULT_GEMINI_PROMPT)
        def load_from_file():
            path = filedialog.askopenfilename(filetypes=[("Текстовые файлы", "*.txt")])
            if path:
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        prompt_text.delete(1.0, tk.END); prompt_text.insert(tk.END, f.read())
                except Exception as e: messagebox.showerror("Ошибка чтения", f"Не удалось прочитать файл: {e}")
        ttk.Button(top_panel, text="Использовать промт по умолчанию", command=use_default).pack(side=tk.LEFT)
        ttk.Button(top_panel, text="Загрузить промт...", command=load_from_file).pack(side=tk.LEFT, padx=10)
        def change_api_key():
            new_key = self._custom_ask_string("Смена API ключа", "Введите новый API ключ Gemini:", parent=analyzer_window)
            if new_key:
                self.gemini_api_key = new_key
                self._save_api_key_to_file(new_key)
                status_label.config(text="API ключ обновлен. Можно запускать анализ.")
        ttk.Button(top_panel, text="Сменить API ключ...", command=change_api_key).pack(side=tk.RIGHT)
        prompt_frame = ttk.Labelframe(main_frame, text="Промт для Gemini (можно редактировать)")
        prompt_frame.pack(fill=tk.BOTH, expand=True)
        prompt_text = scrolledtext.ScrolledText(prompt_frame, wrap=tk.WORD, height=10, font=("Segoe UI", 10))
        prompt_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        use_default()
        bottom_panel = ttk.Frame(main_frame)
        bottom_panel.pack(fill=tk.X, pady=(10, 0))
        def start_analysis():
            user_prompt = prompt_text.get(1.0, tk.END).strip()
            if not user_prompt:
                messagebox.showwarning("Пустой промт", "Пожалуйста, введите промт.", parent=analyzer_window)
                return
            run_button.config(state="disabled")
            status_label.config(text="Проверка ключа и запуск анализа...")
            threading.Thread(target=self.run_gemini_analysis_thread, args=(user_prompt, self.last_clean_transcription, analyzer_window, status_label, run_button), daemon=True).start()
        run_button = ttk.Button(bottom_panel, text="Запустить анализ", command=start_analysis)
        run_button.pack()
        status_label = ttk.Label(bottom_panel, text="Ключ загружен из файла, если он был найден.")
        status_label.pack(pady=5)
    def run_gemini_analysis_thread(self, prompt, text_to_analyze, window, status_label, button):
        try:
            if not self.gemini_api_key:
                self.log("API ключ не найден, запрашиваю у пользователя...", tag='INFO')
                self.root.after(0, self.request_and_save_new_key, window)
                status_label.config(text="Введите API ключ в появившемся диалоге.")
                button.config(state="normal")
                return
            self.log(f"Начинаю анализ с Gemini. Ключ: ...{self.gemini_api_key[-4:]}", tag='INFO')
            genai.configure(api_key=self.gemini_api_key)
            model = genai.GenerativeModel('gemini-1.5-flash-latest')
            final_prompt = f"{prompt}\n\n{text_to_analyze}"
            response = model.generate_content(final_prompt)
            self.log("Анализ Gemini успешно завершен.", tag='SUCCESS')
            self.root.after(0, window.destroy)
            self.root.after(0, self.show_gemini_result_window, response.text)
        except Exception as e:
            error_message = str(e)
            self.log(f"Ошибка при обращении к API Gemini: {error_message}", tag='ERROR')
            if "API_KEY_INVALID" in error_message or "permission_denied" in error_message.lower():
                 messagebox.showerror("Ошибка API", f"Ваш API ключ недействителен или не имеет доступа.\n\nПожалуйста, введите корректный ключ.", parent=window)
                 self.gemini_api_key = None
                 status_label.config(text="Ошибка: API ключ недействителен. Попробуйте сменить его.")
            else:
                 messagebox.showerror("Ошибка API", f"Произошла ошибка при анализе:\n{error_message}", parent=window)
                 status_label.config(text=f"Ошибка: {error_message}")
            button.config(state="normal")
    def request_and_save_new_key(self, parent_window):
        new_key = self._custom_ask_string("API Ключ Gemini", "API ключ не найден или недействителен.\nПожалуйста, введите ваш ключ:", parent=parent_window)
        if new_key:
            self.gemini_api_key = new_key
            self._save_api_key_to_file(new_key)
            messagebox.showinfo("Ключ сохранен", "Ключ успешно сохранен. Теперь можно запустить анализ.", parent=parent_window)
    def show_gemini_result_window(self, result_text):
        result_window = Toplevel(self.root)
        result_window.title("Результат анализа Gemini")
        result_window.geometry("800x650")
        result_frame = ttk.Frame(result_window, padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True)
        text_area = scrolledtext.ScrolledText(result_frame, wrap=tk.WORD, font=("Segoe UI", 10))
        text_area.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        text_area.insert(tk.END, result_text)
        text_area.bind("<Key>", lambda e: "break")
        text_area.bind("<Button-3>", self._show_context_menu)
        buttons_frame = ttk.Frame(result_frame)
        buttons_frame.pack(fill=tk.X)
        def save_as_txt(text):
            path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Текстовые файлы", "*.txt")], title="Сохранить как текст")
            if path:
                with open(path, 'w', encoding='utf-8') as f: f.write(text)
                messagebox.showinfo("Сохранено", f"Файл успешно сохранен в {path}", parent=result_window)
        def save_as_docx(text):
            path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Документы Word", "*.docx")], title="Сохранить как Word")
            if not path: return
            doc = Document()
            for line in text.split('\n'):
                line = line.strip()
                if line.startswith("## "): doc.add_heading(line.lstrip("## ").strip(), level=2)
                elif line.startswith("# "): doc.add_heading(line.lstrip("# ").strip(), level=1)
                elif line.startswith("* "):
                    p = doc.add_paragraph(line.lstrip("* ").strip()); p.style = 'List Bullet'
                elif line.startswith("— "):
                    p = doc.add_paragraph(); p.add_run(f"— {line.lstrip('— ').strip()}")
                elif line.startswith("|") and line.endswith("|"): doc.add_paragraph(line, style='No Spacing')
                else: doc.add_paragraph(line)
            try:
                doc.save(path)
                messagebox.showinfo("Сохранено", f"Документ Word успешно сохранен в {path}", parent=result_window)
            except Exception as e: messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить документ: {e}", parent=result_window)
        def copy_text(text):
            result_window.clipboard_clear()
            result_window.clipboard_append(text)
            messagebox.showinfo("Скопировано", "Результат скопирован в буфер обмена.", parent=result_window)
        ttk.Button(buttons_frame, text="Сохранить как Word (.docx)...", command=lambda: save_as_docx(result_text)).pack(side=tk.LEFT)
        ttk.Button(buttons_frame, text="Сохранить как Текст (.txt)...", command=lambda: save_as_txt(result_text)).pack(side=tk.LEFT, padx=10)
        ttk.Button(buttons_frame, text="Копировать в буфер", command=lambda: copy_text(result_text)).pack(side=tk.LEFT)

# --- Точка входа в программу ---
if __name__ == "__main__":
    # Патчи для скрытия окон должны быть применены до создания App
    root = tk.Tk()
    app = WhisperApp(root)
    if root.winfo_exists():
        root.mainloop()
